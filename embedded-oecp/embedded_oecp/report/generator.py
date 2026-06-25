# **********************************************************************************
# Copyright (c) Huawei Technologies Co., Ltd. 2026-2026. All rights reserved.
# [embedded-oecp] is licensed under the Mulan PSL v2.
# You can use this software according to the terms and conditions of the Mulan PSL v2.
# You may obtain a copy of Mulan PSL v2 at:
#     http://license.coscl.org.cn/MulanPSL2
# THIS SOFTWARE IS PROVIDED ON AN "AS IS" BASIS, WITHOUT WARRANTIES OF ANY KIND, EITHER EXPRESS OR
# IMPLIED, INCLUDING BUT NOT LIMITED TO NON-INFRINGEMENT, MERCHANTABILITY OR FIT FOR A PARTICULAR
# PURPOSE.
# See the Mulan PSL v2 for more details.
# Author: lixinyu
# Create: 2025-01-01
# Description: embedded-oecp utility
# **********************************************************************************
import os
import copy
import datetime
from typing import List
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from embedded_oecp.models import TestResult, TestStatus
from embedded_oecp.utils.logger import get_logger

TEMPLATE_MAP = {
    "kernel": {"table_idx": 2, "section": "3.1"},
    "middleware": {"table_idx": 3, "section": "3.2"},
    "package": {"table_idx": 4, "section": "3.3"},
    "compiler": {"table_idx": 5, "section": "4.1"},
    "project": {"table_idx": 6, "section": "4.2"},
    "pkglist": {"table_idx": 7, "section": "4.3"},
    "libc_runtime": {"table_idx": 8, "section": "5.1"},
    "at_test": {"table_idx": 9, "section": "5.2"},
    "posix": {"table_idx": 10, "section": "5.3"},
}

SUMMARY_TABLE_ROW_MAP = {
    "kernel": 1,
    "middleware": 2,
    "package": 3,
    "compiler": 4,
    "project": 5,
    "pkglist": 6,
    "libc_runtime": 7,
    "at_test": 8,
    "posix": 9,
}


def _docx_element(obj):
    return getattr(obj, "_element")


def _row_tr(row):
    return getattr(row, "_tr")


def _table_tbl(table):
    return getattr(table, "_tbl")


class ReportGenerator:
    def __init__(self, output_dir: str, template_path: str = None):
        self.output_dir = output_dir
        self._ole_counter = 0
        if template_path:
            self.template_path = template_path
        else:
            pkg_template = os.path.join(os.path.dirname(__file__), "..", "assets", "report_template.docx")
            cwd_template = os.path.join(os.getcwd(), "openEuler Embedded OSV 兼容性认证测试报告模板_V1.0.docx")
            if os.path.isfile(pkg_template):
                self.template_path = pkg_template
            elif os.path.isfile(cwd_template):
                self.template_path = cwd_template
            else:
                self.template_path = cwd_template
        os.makedirs(output_dir, exist_ok=True)

    def generate(self, results: List[TestResult], conclusion: str = "", config: dict = None):
        logger = get_logger()

        if not os.path.isfile(self.template_path):
            logger.warning(f"模板文件不存在: {self.template_path}，仅生成 JSON/TXT 报告")
            self._generate_txt_json(results, conclusion)
            return

        doc = Document(self.template_path)

        self._remove_security_section(doc)

        self._insert_cover_page(doc, config)
        self._adjust_table_widths(doc)
        self._fill_basic_info(doc, config)
        self._fill_test_results(doc, results)
        self._fill_summary_table(doc, results)
        self._fill_conclusion(doc, conclusion)
        self._set_table_fonts(doc)

        ts = datetime.datetime.now(datetime.timezone.utc).strftime("%Y%m%d_%H%M%S")
        docx_path = os.path.join(self.output_dir, f"openEuler_Embedded_OSV_认证测试报告_{ts}.docx")
        doc.save(docx_path)
        logger.info(f"DOCX report: {docx_path}")

        evidence_dir = os.path.join(self.output_dir, "evidence")
        self._embed_excel_attachments(docx_path, evidence_dir)

        self._generate_txt_json(results, conclusion)

    @staticmethod
    def _remove_security_section(doc: Document):
        logger = get_logger()
        security_table = None
        for table in doc.tables:
            if _docx_element(table).getparent() is not None:
                for row in table.rows:
                    for cell in row.cells:
                        if "安全配置" in cell.text or "安全认证" in cell.text:
                            security_table = table
                            break
                    if security_table:
                        break
            if security_table:
                break

        if security_table is None:
            return

        _docx_element(security_table).getparent().remove(_docx_element(security_table))
        logger.info("Removed security table from template")

        body = doc.element.body
        to_remove = []
        for child in body:
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if tag == "p":
                text = "".join(child.itertext()).strip()
                if "六、安全认证" in text or "6.1 安全配置检查" in text:
                    to_remove.append(child)

        for elem in to_remove:
            body.remove(elem)
            logger.info(f"Removed paragraph from template")

        renumber_map = {
            "七、测试结果汇总": "六、测试结果汇总",
            "八、认证结论": "七、认证结论",
            "九、附录": "八、附录",
        }
        for child in body:
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if tag == "p":
                text = "".join(child.itertext()).strip()
                for old_text, new_text in renumber_map.items():
                    if old_text in text:
                        for run_elem in child.iter(qn("w:t")):
                            if old_text in (run_elem.text or ""):
                                run_elem.text = run_elem.text.replace(old_text, new_text)
                        break

        summary_table = None
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip() == "安全配置" or cell.text.strip() == "安全配置检查":
                        summary_table = table
                        row_to_remove = row
                        break
                if summary_table:
                    break
            if summary_table:
                break

        if summary_table and row_to_remove:
            _docx_element(summary_table).remove(_row_tr(row_to_remove))
            logger.info("Removed security row from summary table")

        env_table = None
        for table in doc.tables:
            for row in table.rows:
                if "构建容器版本" in row.cells[0].text:
                    env_table = table
                    _row_tr(row).getparent().remove(_row_tr(row))
                    logger.info("Removed 构建容器版本 row from env table")
                    break
            if env_table:
                break

    @staticmethod
    def _insert_cover_page(doc: Document, config: dict):
        from docx.oxml.ns import qn as _qn
        from lxml import etree

        body = doc.element.body

        system_name = "XXX-OS"
        company_name = config.get("company", "xxxx公司") if config else "xxxx公司"

        cover_elements = []

        for _ in range(9):
            p = doc.add_paragraph()
            cover_elements.append(_docx_element(p))

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(system_name)
        run.font.size = Pt(42)
        run.font.color.rgb = RGBColor(0x00, 0xB0, 0xF0)
        run.font.name = "宋体"
        _docx_element(run).rPr.rFonts.set(_qn("w:eastAsia"), "宋体")
        cover_elements.append(_docx_element(p))

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("兼容性认证测试报告")
        run.font.size = Pt(42)
        run.font.name = "宋体"
        _docx_element(run).rPr.rFonts.set(_qn("w:eastAsia"), "宋体")
        cover_elements.append(_docx_element(p))

        for _ in range(3):
            p = doc.add_paragraph()
            cover_elements.append(_docx_element(p))

        for label in ["拟  制", "审  核"]:
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = Pt(18)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(label)
            run.font.name = "宋体"
            _docx_element(run).rPr.rFonts.set(_qn("w:eastAsia"), "宋体")
            cover_elements.append(_docx_element(p))

        for _ in range(2):
            p = doc.add_paragraph()
            cover_elements.append(_docx_element(p))

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(company_name)
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x00, 0xB0, 0xF0)
        run.font.name = "宋体"
        _docx_element(run).rPr.rFonts.set(_qn("w:eastAsia"), "宋体")
        cover_elements.append(_docx_element(p))

        for _ in range(2):
            p = doc.add_paragraph()
            cover_elements.append(_docx_element(p))

        p = doc.add_paragraph()
        p_pr = _docx_element(p).get_or_add_pPr()
        sect_pr = etree.SubElement(p_pr, _qn("w:sectPr"))
        pg_sz = etree.SubElement(sect_pr, _qn("w:pgSz"))
        pg_sz.set(_qn("w:w"), "11906")
        pg_sz.set(_qn("w:h"), "16838")
        pg_mar = etree.SubElement(sect_pr, _qn("w:pgMar"))
        pg_mar.set(_qn("w:top"), "1440")
        pg_mar.set(_qn("w:right"), "1800")
        pg_mar.set(_qn("w:bottom"), "1440")
        pg_mar.set(_qn("w:left"), "1800")
        pg_mar.set(_qn("w:header"), "851")
        pg_mar.set(_qn("w:footer"), "992")
        pg_mar.set(_qn("w:gutter"), "0")
        sect_pr.set(_qn("w:type"), "nextPage")
        cover_elements.append(_docx_element(p))

        for elem in reversed(cover_elements):
            body.remove(elem)
            body.insert(0, elem)

    @staticmethod
    def _adjust_table_widths(doc: Document):
        for ti in range(2, 11):
            if ti >= len(doc.tables):
                continue
            table = doc.tables[ti]
            for row in table.rows:
                row.cells[0].width = Cm(3.5)
                row.cells[1].width = Cm(13.5)
        for ti in range(0, 2):
            if ti >= len(doc.tables):
                continue
            table = doc.tables[ti]
            for row in table.rows:
                row.cells[0].width = Cm(4.0)
                row.cells[1].width = Cm(13.0)
                row.cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                row.cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        if len(doc.tables) > 11:
            summary_table = doc.tables[11]
            summary_col_widths = [Cm(1.8), Cm(2.2), Cm(4.5), Cm(6.0), Cm(2.5)]
            for row in summary_table.rows:
                for ci, w in enumerate(summary_col_widths):
                    if ci < len(row.cells):
                        row.cells[ci].width = w
        if len(doc.tables) > 12:
            conclusion_table = doc.tables[12]
            for row in conclusion_table.rows:
                row.cells[0].width = Cm(3.5)
                row.cells[1].width = Cm(13.5)

    @staticmethod
    def _set_table_fonts(doc: Document):
        from docx.oxml.ns import qn as _qn
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.name = "宋体"
                            run.font.size = Pt(10.5)
                            _docx_element(run).rPr.rFonts.set(_qn("w:eastAsia"), "宋体")

    def _fill_basic_info(self, doc: Document, config: dict):
        if not config:
            return
        logger = get_logger()

        if len(doc.tables) < 2:
            return

        info_table = doc.tables[0]
        env_table = doc.tables[1]

        replacements_info = {}
        basic_defaults = {
            0: "（自行填写）",
            1: "24.03-LTS",
            3: "rk3561",
            4: "（自行填写）",
            5: "□铁路 □金融 □电信 □水利 □教育 □政府 □安平\n□广电 □卫生 □制造 □能源 □民航 □公路水运\n□运营商",
            6: "（自行填写）",
        }
        replacements_info.update(basic_defaults)
        if config.get("arch"):
            replacements_info[2] = str(config["arch"])
        ts = datetime.datetime.now(datetime.timezone.utc).strftime("%Y-%m-%d")
        replacements_info[7] = ts

        for row_idx, value in replacements_info.items():
            if row_idx < len(info_table.rows):
                cell = info_table.rows[row_idx].cells[1]
                cell.text = value

        env_table.rows[0].cells[1].text = "（自行填写）"
        env_table.rows[1].cells[1].text = "（自行填写）"

        build_ts = config.get("build_timestamp", "")
        if build_ts:
            env_table.rows[3].cells[1].text = build_ts

        env_table.rows[4].cells[1].text = "（自行填写）"

        sdk_file = config.get("sdk_file", "")
        if sdk_file and os.path.isfile(sdk_file):
            import hashlib
            md5 = hashlib.md5()
            with open(sdk_file, "rb") as f:
                for chunk in iter(lambda: f.read(8192), b""):
                    md5.update(chunk)
            sdk_name = os.path.basename(sdk_file)
            env_table.rows[5].cells[1].text = f"{sdk_name}\nMD5: {md5.hexdigest()}"
        elif sdk_file:
            env_table.rows[5].cells[1].text = os.path.basename(sdk_file)

        self._insert_os_release_row(doc, env_table, config)

    def _insert_os_release_row(self, doc: Document, env_table, config: dict):
        logger = get_logger()
        device_cfg = config.get("device", {})
        device_ip = device_cfg.get("ip")
        if not device_ip:
            return

        from embedded_oecp.utils.shell import run_remote_command
        from embedded_oecp.utils.terminal_screenshot import render_terminal_screenshot
        from docx.oxml.ns import qn as _qn

        host = device_ip
        user = device_cfg.get("user", "root")
        password = device_cfg.get("password", "")
        port = device_cfg.get("port", 22)

        stdout, stderr, rc = run_remote_command(
            host, user, password, "cat /etc/os-release", port=port,
        )

        version_text = ""
        full_output = ""
        if rc == 0 and stdout.strip():
            full_output = stdout.strip()
            for line in full_output.split("\n"):
                if line.startswith("PRETTY_NAME="):
                    version_text = line.split("=", 1)[1].strip('"')
                    break
            if not version_text:
                for line in full_output.split("\n"):
                    if line.startswith("NAME=") or line.startswith("VERSION="):
                        version_text += line.split("=", 1)[1].strip('"') + " "
                version_text = version_text.strip()
        if not version_text:
            version_text = full_output.split("\n")[0] if full_output else "未知"

        logger.info(f"Device OS release: {version_text}")

        evidence_dir = os.path.join(self.output_dir, "evidence")
        os.makedirs(evidence_dir, exist_ok=True)

        lines = [{"text": "$ cat /etc/os-release", "type": "command"}]
        if stdout:
            for line in stdout.strip().split("\n"):
                lines.append({"text": line, "color": "#cdd6f4"})
        screenshot_path = os.path.join(evidence_dir, "os_release.png")
        render_terminal_screenshot("被测操作系统版本", lines, screenshot_path)

        tbl = _table_tbl(env_table)
        test_tool_row_idx = None
        for i, row in enumerate(env_table.rows):
            if "测试工具" in row.cells[0].text:
                test_tool_row_idx = i
                break

        if test_tool_row_idx is None:
            return

        new_row_xml = copy.deepcopy(_row_tr(env_table.rows[test_tool_row_idx]))
        _row_tr(env_table.rows[test_tool_row_idx]).addnext(new_row_xml)

        new_row = None
        for row in env_table.rows:
            if _row_tr(row) is new_row_xml:
                new_row = row
                break

        if new_row is None:
            return

        new_row.cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        new_row.cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        new_row.cells[0].paragraphs[0].clear()
        run_label = new_row.cells[0].paragraphs[0].add_run("被测操作系统版本")
        run_label.font.size = Pt(9)
        run_label.font.bold = True

        cell = new_row.cells[1]
        cell.text = ""

        first_para = cell.paragraphs[0]
        if full_output:
            for li, line in enumerate(full_output.split("\n")):
                if li == 0:
                    para = first_para
                else:
                    para = cell.add_paragraph()
                run = para.add_run(line)
                run.font.size = Pt(8)
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(0)
                para.paragraph_format.line_spacing = Pt(12)

        img_para = cell.add_paragraph()
        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img_para.paragraph_format.space_before = Pt(4)
        img_para.paragraph_format.space_after = Pt(4)
        img_run = img_para.add_run()
        try:
            img_run.add_picture(screenshot_path, width=Cm(13.0))
        except Exception:
            try:
                img_run.add_picture(screenshot_path, width=Inches(5.0))
            except Exception as e:
                logger.debug(f"Failed to add OS release picture: {e}")

        logger.info(f"Inserted OS release row: {version_text}")

    def _fill_test_results(self, doc: Document, results: List[TestResult]):
        logger = get_logger()

        tested_checkers = set()
        for r in results:
            status = r.status if isinstance(r.status, str) else r.status
            if status != TestStatus.NONE.value:
                checker_name = self._get_checker_name(r)
                if checker_name:
                    tested_checkers.add(checker_name)

        for checker_name in tested_checkers:
            info = TEMPLATE_MAP.get(checker_name)
            if not info:
                continue
            table_idx = info["table_idx"]
            if table_idx >= len(doc.tables):
                continue

            table = doc.tables[table_idx]
            checker_results = [r for r in results if self._get_checker_name(r) == checker_name]

            if not checker_results:
                continue

            row_idx = self._find_result_row(table)
            if row_idx is None:
                continue

            cell = table.rows[row_idx].cells[1]
            cell.text = ""

            evidence_dir = os.path.join(self.output_dir, "evidence")
            evi_files = self._find_evidence_files(checker_name, evidence_dir)

            step_labels = self._build_step_labels(checker_name, len(checker_results))

            first_para = True
            for i, r in enumerate(checker_results):
                d = r.to_dict() if hasattr(r, "to_dict") else r
                status = d.get("status", "")
                detail = d.get("detail", "")
                label = step_labels[i] if i < len(step_labels) else f"Step {i+1}"

                if first_para:
                    para = cell.paragraphs[0]
                    first_para = False
                else:
                    para = cell.add_paragraph()

                run_status = para.add_run(f"[{status}] ")
                run_status.font.size = Pt(9)
                run_status.font.bold = True
                if status == "PASS":
                    run_status.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
                elif status == "FAIL":
                    run_status.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

                run_detail = para.add_run(f"{label}: ")
                run_detail.font.size = Pt(9)

                if checker_name == "kernel" and i == 0:
                    for part in detail.split(", "):
                        para = cell.add_paragraph()
                        run_part = para.add_run(part)
                        run_part.font.size = Pt(9)
                else:
                    run_detail2 = para.add_run(detail)
                    run_detail2.font.size = Pt(9)

                evi_file_list = self._find_step_evidence(checker_name, i, evi_files)

                if checker_name in ("package", "pkglist"):
                    self._insert_excel_ole(cell, evidence_dir)
                else:
                    for evi_file in evi_file_list:
                        img_path = os.path.join(evidence_dir, evi_file)
                        img_para = cell.add_paragraph()
                        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        img_para.paragraph_format.space_before = Pt(4)
                        img_para.paragraph_format.space_after = Pt(4)
                        img_run = img_para.add_run()
                        try:
                            img_width = Cm(13.0)
                            img_run.add_picture(img_path, width=img_width)
                        except Exception:
                            try:
                                from PIL import Image as PILImage
                                with PILImage.open(img_path) as im:
                                    img_w, img_h = im.size
                                ratio = 13.0 / (img_w / 96.0 * 2.54)
                                new_w = Cm(13.0)
                                new_h = Cm(img_h / 96.0 * 2.54 * ratio)
                                img_run.add_picture(img_path, width=new_w, height=new_h)
                            except Exception:
                                img_run.add_picture(img_path, width=Inches(5.0))

            if checker_name in ("at_test", "posix"):
                log_names = {
                    "at_test": ("at_test_execute_log.txt", "AT"),
                    "posix": ("posix_execute_log.txt", "POSIX"),
                }
                log_info = log_names.get(checker_name, ("", ""))
                log_file, test_label = log_info
                log_path = os.path.join(evidence_dir, log_file)
                if os.path.isfile(log_path):
                    self._insert_log_attachment(cell, log_path, test_label)

        for checker_name, info in TEMPLATE_MAP.items():
            if checker_name not in tested_checkers:
                table_idx = info["table_idx"]
                if table_idx < len(doc.tables):
                    table = doc.tables[table_idx]
                    row_idx = self._find_result_row(table)
                    if row_idx is not None:
                        table.rows[row_idx].cells[1].text = "未测试"

    def _fill_summary_table(self, doc: Document, results: List[TestResult]):
        summary_table_idx = 11
        if summary_table_idx >= len(doc.tables):
            return

        table = doc.tables[summary_table_idx]

        merged = {}
        for r in results:
            checker_name = self._get_checker_name(r)
            if checker_name not in merged:
                merged[checker_name] = []
            merged[checker_name].append(r)

        for checker_name, checker_results in merged.items():
            row_idx = SUMMARY_TABLE_ROW_MAP.get(checker_name)
            if row_idx is None or row_idx >= len(table.rows):
                continue

            statuses = []
            for r in checker_results:
                d = r.to_dict() if hasattr(r, "to_dict") else r
                statuses.append(d.get("status", "NONE"))

            if any(s == "FAIL" for s in statuses):
                final_status = "FAIL"
            elif all(s == "PASS" for s in statuses):
                final_status = "PASS"
            elif any(s == "PASS" for s in statuses):
                final_status = "PASS"
            else:
                final_status = "NONE"

            last_col_idx = len(table.columns) - 1
            cell = table.rows[row_idx].cells[last_col_idx]
            cell.text = ""
            run = cell.paragraphs[0].add_run(final_status)
            run.font.size = Pt(10)
            run.font.bold = True
            if final_status == "PASS":
                run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
            elif final_status == "FAIL":
                run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

        for checker_name, row_idx in SUMMARY_TABLE_ROW_MAP.items():
            if checker_name not in merged:
                if row_idx < len(table.rows):
                    last_col_idx = len(table.columns) - 1
                    cell = table.rows[row_idx].cells[last_col_idx]
                    cell.text = ""
                    run = cell.paragraphs[0].add_run("NONE")
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    def _fill_conclusion(self, doc: Document, conclusion: str):
        conclusion_table_idx = 12
        if conclusion_table_idx >= len(doc.tables):
            return

        table = doc.tables[conclusion_table_idx]

        if "通过" in conclusion and "不" not in conclusion:
            label = "☑通过  □不通过  □待整改"
            color = RGBColor(0x00, 0x80, 0x00)
        elif "待整改" in conclusion:
            label = "□通过  □不通过  ☑待整改"
            color = RGBColor(0xCC, 0x99, 0x00)
        else:
            label = "□通过  ☑不通过  □待整改"
            color = RGBColor(0xCC, 0x00, 0x00)

        cell = table.rows[0].cells[1]
        cell.text = ""
        run = cell.paragraphs[0].add_run(label)
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = color

        fail_results = []
        for ti in range(2, 11):
            if ti < len(doc.tables):
                t = doc.tables[ti]
                rr = self._find_result_row(t)
                if rr is not None:
                    txt = t.rows[rr].cells[1].text
                    if "FAIL" in txt:
                        test_name = t.rows[0].cells[1].text[:30]
                        fail_results.append(test_name)

        if fail_results:
            table.rows[1].cells[1].text = "需整改项: " + "; ".join(fail_results)

    @staticmethod
    def _find_evidence_files(checker_name: str, evidence_dir: str) -> list:
        if not os.path.isdir(evidence_dir):
            return []
        prefix_map = {
            "kernel": "kernel_step",
            "middleware": "middleware_step",
            "package": "package_step",
            "compiler": "compiler_step",
            "project": "project_step",
            "pkglist": "pkglist_step",
            "libc_runtime": "libc_runtime_step",
            "at_test": "at_test_step",
            "posix": "posix_step",
        }
        prefix = prefix_map.get(checker_name, "")
        if not prefix:
            return []
        return sorted([f for f in os.listdir(evidence_dir) if f.startswith(prefix) and f.endswith((".png", ".jpg"))])

    @staticmethod
    def _build_step_labels(checker_name: str, count: int) -> list:
        labels_map = {
            "kernel": [
                "Step1-获取manifest信息",
                "Step2-内核仓库在openEuler组织",
                "Step3-本地内核仓与manifest一致",
                "Step4-内核版本检查",
                "Step5-目标设备内核版本",
            ],
            "middleware": [
                "Step1-目标设备libc.so MD5提取",
                "Step2-env.yaml基线获取",
                "Step3-libc.so MD5比对",
            ],
            "package": [
                "Step1-镜像依赖包统计(详见Excel)",
            ],
            "compiler": [
                "Step1-编译工具链基线获取",
                "Step2-gcc和glibc MD5提取",
                "Step3-gcc和glibc MD5比对结论",
            ],
            "project": [
                "Step1-构建命令检查",
                "Step2-构建目录结构检查",
            ],
            "pkglist": [
                "Step1-社区包版本一致性检查(详见Excel)",
            ],
            "libc_runtime": [
                "Step1-上传检测程序并校验MD5",
                "Step2-运行检测程序验证通过",
            ],
            "at_test": [
                "Step1-执行AT测试",
                "Step2-AT测试结果汇总",
            ],
            "posix": [
                "Step1-执行POSIX测试",
                "Step2-POSIX测试结果汇总",
            ],
        }
        return labels_map.get(checker_name, [f"Step {i+1}" for i in range(count)])

    @staticmethod
    def _find_step_evidence(checker_name: str, step_idx: int, evi_files: list) -> list:
        step_key_map = {
            "kernel": [None, "step2", "step3", "step4", "step5"],
            "middleware": ["step1", "step2", "step3"],
            "package": ["step1"],
            "compiler": ["step1", "step2", "step3"],
            "project": ["step1", "step2"],
            "pkglist": ["step1"],
            "libc_runtime": [["step1a", "step1b"], ["step2"]],
            "at_test": ["step1", "step2"],
            "posix": ["step1", "step2"],
        }

        keys = step_key_map.get(checker_name)
        if keys:
            if step_idx < len(keys):
                key = keys[step_idx]
                if key is None:
                    return []
                if isinstance(key, list):
                    matched = []
                    for k in key:
                        matched.extend(f for f in evi_files if k in f)
                    return matched
                matched = [f for f in evi_files if key in f]
                return matched
        else:
            if step_idx < len(evi_files):
                return [evi_files[step_idx]]
        return []

    @staticmethod
    def _find_result_row(table) -> int:
        for i, row in enumerate(table.rows):
            if row.cells[0].text.strip() == "测试结果":
                return i
        return None

    @staticmethod
    def _find_note_row(table) -> int:
        for i, row in enumerate(table.rows):
            if row.cells[0].text.strip() == "备注":
                return i
        return None

    @staticmethod
    def _get_checker_name(result) -> str:
        d = result.to_dict() if hasattr(result, "to_dict") else result
        sub_item = d.get("sub_item", "")
        mapping = {
            "内核": "kernel",
            "基础中间件": "middleware",
            "其他软件包": "package",
            "编译链": "compiler",
            "构建工程": "project",
            "包列表": "pkglist",
            "C库运行时": "libc_runtime",
            "基础功能": "at_test",
            "POSIX": "posix",
        }
        return mapping.get(sub_item, "")

    def _generate_txt_json(self, results: List[TestResult], conclusion: str):
        logger = get_logger()

        import json
        json_path = os.path.join(self.output_dir, "summary.json")
        with open(json_path, "w", encoding="utf-8") as f:
            json.dump({
                "conclusion": conclusion,
                "results": [r.to_dict() if hasattr(r, "to_dict") else r for r in results],
            }, f, ensure_ascii=False, indent=2)
        logger.info(f"JSON report: {json_path}")

        txt_path = os.path.join(self.output_dir, "summary.txt")
        with open(txt_path, "w", encoding="utf-8") as f:
            f.write(f"{'项目':<8} {'子项':<14} {'要求':<20} {'状态':<8} {'说明'}\n")
            f.write("-" * 70 + "\n")
            for r in results:
                d = r.to_dict() if hasattr(r, "to_dict") else r
                txt_line = (
                    f"{d['category']:<8} {d['sub_item']:<14} "
                    f"{d['requirement']:<20} {d['status']:<8} "
                    f"{d.get('detail', '')}\n"
                )
                f.write(txt_line)
            f.write(f"\n认证结论: {conclusion}\n")
        logger.info(f"TXT report: {txt_path}")

    @staticmethod
    def _create_ole_package_bin(file_bytes: bytes, filename: str = "package", filepath: str = "") -> bytes:
        import struct
        import io

        sector_size = 512
        end_of_chain = 0xFFFFFFFE
        free_sector = 0xFFFFFFFF
        no_stream = 0xFFFFFFFF

        wps_clsid = "20a70df22fc0ce11927b0800095ae340"

        objinfo_data = b'@\x00\x03\x00\x01\x00'
        mini_cutoff = 4096

        if not filepath:
            filepath = f"/tmp/embedded-oecp/report/evidence/{filename}"

        fname_bytes = filename.encode('utf-8') + b'\x00'
        fpath_bytes = filepath.encode('utf-8') + b'\x00'

        ole10native = bytearray()
        ole10native += struct.pack('<H', 2)
        ole10native += fname_bytes
        ole10native += fpath_bytes
        ole10native += struct.pack('<H', 0)
        ole10native += struct.pack('<H', 3)
        ole10native += struct.pack('<I', len(fpath_bytes))
        ole10native += fpath_bytes
        ole10native += struct.pack('<I', len(file_bytes))
        ole10native += file_bytes
        ole10native += b'\x00\x00\x00\x00'

        native_size = len(ole10native)
        ole10native_with_header = struct.pack('<I', native_size) + bytes(ole10native)

        pkg_sectors = (len(ole10native_with_header) + sector_size - 1) // sector_size
        num_data_sectors = max(pkg_sectors, 1)

        ministream_sectors = (len(objinfo_data) + sector_size - 1) // sector_size

        entries_per_fat_sector = sector_size // 4

        overhead_sectors = 1 + 1 + 1 + ministream_sectors
        total_sectors = overhead_sectors + num_data_sectors
        num_fat_sectors = (total_sectors + entries_per_fat_sector - 1) // entries_per_fat_sector

        if num_fat_sectors > 1:
            total_sectors = num_fat_sectors + 1 + 1 + 1 + ministream_sectors + num_data_sectors
            num_fat_sectors2 = (total_sectors + entries_per_fat_sector - 1) // entries_per_fat_sector
            if num_fat_sectors2 != num_fat_sectors:
                num_fat_sectors = num_fat_sectors2
                total_sectors = num_fat_sectors + 1 + 1 + 1 + ministream_sectors + num_data_sectors

        layout = {
            "pkg_start": 0,
            "ministream": num_data_sectors,
            "minifat": num_data_sectors + ministream_sectors,
            "dir": num_data_sectors + ministream_sectors + 1,
            "fat_start": num_data_sectors + ministream_sectors + 2,
        }

        total_sectors = layout["fat_start"] + num_fat_sectors
        num_fat_sectors_check = (total_sectors + entries_per_fat_sector - 1) // entries_per_fat_sector
        if num_fat_sectors_check != num_fat_sectors:
            num_fat_sectors = num_fat_sectors_check
            total_sectors = layout["fat_start"] + num_fat_sectors

        header = bytearray(sector_size)
        header[0:8] = b'\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1'
        header[8:24] = bytes.fromhex(wps_clsid)
        struct.pack_into('<H', header, 24, 0x003E)
        struct.pack_into('<H', header, 26, 0x0003)
        struct.pack_into('<H', header, 28, 0xFFFE)
        struct.pack_into('<H', header, 30, 9)
        struct.pack_into('<H', header, 32, 6)
        header[34:40] = b'\x00' * 6
        struct.pack_into('<I', header, 40, 0)
        struct.pack_into('<I', header, 44, num_fat_sectors)
        struct.pack_into('<I', header, 48, layout["dir"])
        struct.pack_into('<I', header, 52, 0)
        struct.pack_into('<I', header, 56, mini_cutoff)
        struct.pack_into('<I', header, 60, layout["minifat"])
        struct.pack_into('<I', header, 64, 1)
        struct.pack_into('<I', header, 68, end_of_chain)
        struct.pack_into('<I', header, 72, 0)
        for i in range(num_fat_sectors):
            struct.pack_into('<I', header, 76 + i * 4, layout["fat_start"] + i)
        for i in range(num_fat_sectors, 109):
            struct.pack_into('<I', header, 76 + i * 4, free_sector)

        fat_sect = 0xFFFFFFFD

        fat = bytearray(num_fat_sectors * sector_size)

        for i in range(num_fat_sectors):
            struct.pack_into('<I', fat, i * entries_per_fat_sector * 4, end_of_chain)

        for i in range(layout["fat_start"], layout["fat_start"] + num_fat_sectors):
            struct.pack_into('<I', fat, i * 4, fat_sect)
        struct.pack_into('<I', fat, layout["dir"] * 4, end_of_chain)
        struct.pack_into('<I', fat, layout["minifat"] * 4, end_of_chain)
        for i in range(ministream_sectors):
            if i < ministream_sectors - 1:
                struct.pack_into('<I', fat, (layout["ministream"] + i) * 4, layout["ministream"] + i + 1)
            else:
                struct.pack_into('<I', fat, (layout["ministream"] + i) * 4, end_of_chain)

        for i in range(num_data_sectors - 1):
            struct.pack_into('<I', fat, (layout["pkg_start"] + i) * 4, layout["pkg_start"] + i + 1)
        struct.pack_into('<I', fat, (layout["pkg_start"] + num_data_sectors - 1) * 4, end_of_chain)

        fat_entries_used = layout["fat_start"] + num_fat_sectors
        for i in range(fat_entries_used, num_fat_sectors * entries_per_fat_sector):
            struct.pack_into('<I', fat, i * 4, free_sector)

        def direntry(name, obj_type, **fields):
            left = fields.get("left", no_stream)
            right = fields.get("right", no_stream)
            child = fields.get("child", no_stream)
            color = fields.get("color", 1)
            start = fields.get("start", end_of_chain)
            size = fields.get("size", 0)
            clsid = fields.get("clsid")
            name_u16 = name.encode('utf-16-le') + b'\x00\x00'
            e = bytearray(128)
            e[0:len(name_u16)] = name_u16[:64]
            struct.pack_into('<H', e, 64, len(name_u16))
            e[66] = obj_type
            e[67] = color
            struct.pack_into('<I', e, 68, left)
            struct.pack_into('<I', e, 72, right)
            struct.pack_into('<I', e, 76, child)
            if clsid:
                e[80:96] = bytes.fromhex(clsid)
            struct.pack_into('<I', e, 116, start)
            struct.pack_into('<I', e, 120, size)
            return bytes(e)

        root_entry = direntry("Root Entry", 5, child=1, color=0, start=layout["ministream"], size=64, clsid=wps_clsid)
        ole10native_entry = direntry(
            "\x01Ole10Native", 2, left=2, color=1,
            start=layout["pkg_start"],
            size=len(ole10native_with_header),
        )
        objinfo_entry = direntry("\x03ObjInfo", 2, color=0, start=0, size=len(objinfo_data))
        empty_entry = bytes(128)

        dir_sector = bytearray(sector_size)
        dir_sector[0:128] = root_entry
        dir_sector[128:256] = ole10native_entry
        dir_sector[256:384] = objinfo_entry
        dir_sector[384:512] = empty_entry

        minifat = bytearray(sector_size)
        for i in range(sector_size // 4):
            struct.pack_into('<I', minifat, i * 4, free_sector)
        struct.pack_into('<I', minifat, 0, end_of_chain)

        ministream = bytearray(sector_size * ministream_sectors)
        ministream[0:len(objinfo_data)] = objinfo_data

        pkg_data = bytearray(num_data_sectors * sector_size)
        pkg_data[0:len(ole10native_with_header)] = ole10native_with_header

        buf = io.BytesIO()
        buf.write(header)
        buf.write(pkg_data)
        buf.write(ministream)
        buf.write(minifat)
        buf.write(dir_sector)
        for i in range(num_fat_sectors):
            buf.write(fat[i * sector_size:(i + 1) * sector_size])

        return buf.getvalue()

    @staticmethod
    def _create_ole_raw_bin(file_bytes: bytes) -> bytes:
        import struct
        import io

        sector_size = 512
        end_of_chain = 0xFFFFFFFE
        free_sector = 0xFFFFFFFF
        no_stream = 0xFFFFFFFF

        objinfo_data = b'@\x00\x03\x00\x01\x00'
        mini_cutoff = 4096

        pkg_sectors = (len(file_bytes) + sector_size - 1) // sector_size if file_bytes else 1
        num_data_sectors = pkg_sectors
        ministream_sectors = 1
        entries_per_fat_sector = sector_size // 4
        num_fat_sectors = 1
        total_check = 1 + 1 + 1 + ministream_sectors + num_data_sectors
        if total_check > entries_per_fat_sector:
            num_fat_sectors = (total_check + entries_per_fat_sector - 1) // entries_per_fat_sector

        if num_fat_sectors > 1:
            total_check = num_fat_sectors + 1 + 1 + ministream_sectors + num_data_sectors
            num_fat_sectors2 = (total_check + entries_per_fat_sector - 1) // entries_per_fat_sector
            if num_fat_sectors2 != num_fat_sectors:
                num_fat_sectors = num_fat_sectors2

        layout = {
            "pkg_start": 0,
            "ministream": num_data_sectors,
            "minifat": num_data_sectors + ministream_sectors,
            "dir": num_data_sectors + ministream_sectors + 1,
            "fat_start": num_data_sectors + ministream_sectors + 2,
        }

        total_sectors = layout["fat_start"] + num_fat_sectors
        num_fat_sectors_check = (total_sectors + entries_per_fat_sector - 1) // entries_per_fat_sector
        if num_fat_sectors_check != num_fat_sectors:
            num_fat_sectors = num_fat_sectors_check
            total_sectors = layout["fat_start"] + num_fat_sectors

        header = bytearray(sector_size)
        header[0:8] = b'\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1'
        header[8:24] = bytes.fromhex('3008020000000000c000000000000046')
        struct.pack_into('<H', header, 24, 0x003E)
        struct.pack_into('<H', header, 26, 0x0003)
        struct.pack_into('<H', header, 28, 0xFFFE)
        struct.pack_into('<H', header, 30, 9)
        struct.pack_into('<H', header, 32, 6)
        struct.pack_into('<I', header, 44, num_fat_sectors)
        struct.pack_into('<I', header, 48, layout["dir"])
        struct.pack_into('<I', header, 56, mini_cutoff)
        struct.pack_into('<I', header, 60, layout["minifat"])
        struct.pack_into('<I', header, 64, 1)
        struct.pack_into('<I', header, 68, end_of_chain)
        struct.pack_into('<I', header, 72, 0)
        for i in range(num_fat_sectors):
            struct.pack_into('<I', header, 76 + i * 4, layout["fat_start"] + i)
        for i in range(num_fat_sectors, 109):
            struct.pack_into('<I', header, 76 + i * 4, free_sector)

        fat_sect = 0xFFFFFFFD

        fat = bytearray(num_fat_sectors * sector_size)
        for i in range(num_fat_sectors):
            struct.pack_into('<I', fat, i * entries_per_fat_sector * 4, end_of_chain)
        for i in range(num_fat_sectors):
            struct.pack_into('<I', fat, (layout["fat_start"] + i) * 4, fat_sect)
        struct.pack_into('<I', fat, layout["dir"] * 4, end_of_chain)
        struct.pack_into('<I', fat, layout["minifat"] * 4, end_of_chain)
        struct.pack_into('<I', fat, layout["ministream"] * 4, end_of_chain)
        for i in range(num_data_sectors - 1):
            struct.pack_into('<I', fat, (layout["pkg_start"] + i) * 4, layout["pkg_start"] + i + 1)
        struct.pack_into('<I', fat, (layout["pkg_start"] + num_data_sectors - 1) * 4, end_of_chain)

        fat_entries_used = max(layout["dir"] + 1, layout["fat_start"] + num_fat_sectors)
        for i in range(fat_entries_used, num_fat_sectors * entries_per_fat_sector):
            struct.pack_into('<I', fat, i * 4, free_sector)

        def direntry(name, obj_type, **fields):
            left = fields.get("left", no_stream)
            right = fields.get("right", no_stream)
            child = fields.get("child", no_stream)
            color = fields.get("color", 1)
            start = fields.get("start", end_of_chain)
            size = fields.get("size", 0)
            clsid = fields.get("clsid")
            name_u16 = name.encode('utf-16-le') + b'\x00\x00'
            e = bytearray(128)
            e[0:len(name_u16)] = name_u16[:64]
            struct.pack_into('<H', e, 64, len(name_u16))
            e[66] = obj_type
            e[67] = color
            struct.pack_into('<I', e, 68, left)
            struct.pack_into('<I', e, 72, right)
            struct.pack_into('<I', e, 76, child)
            if clsid:
                e[80:96] = bytes.fromhex(clsid)
            struct.pack_into('<I', e, 116, start)
            struct.pack_into('<I', e, 120, size)
            return bytes(e)

        root = direntry(
            "Root Entry", 5, child=1, color=0,
            start=layout["ministream"], size=64,
            clsid="3008020000000000c000000000000046",
        )
        objinfo = direntry("\x03ObjInfo", 2, left=2, color=1, start=0, size=len(objinfo_data))
        package = direntry("package", 2, color=0, start=layout["pkg_start"], size=len(file_bytes))
        empty = bytes(128)

        dir_sector = bytearray(sector_size)
        dir_sector[0:128] = root
        dir_sector[128:256] = objinfo
        dir_sector[256:384] = package
        dir_sector[384:512] = empty

        minifat = bytearray(sector_size)
        for i in range(sector_size // 4):
            struct.pack_into('<I', minifat, i * 4, free_sector)
        struct.pack_into('<I', minifat, 0, end_of_chain)

        ministream = bytearray(sector_size)
        ministream[0:len(objinfo_data)] = objinfo_data

        pkg_data = bytearray(num_data_sectors * sector_size)
        pkg_data[0:len(file_bytes)] = file_bytes

        buf = io.BytesIO()
        buf.write(header)
        buf.write(pkg_data)
        buf.write(ministream)
        buf.write(minifat)
        buf.write(dir_sector)
        for i in range(num_fat_sectors):
            buf.write(fat[i * sector_size:(i + 1) * sector_size])
        return buf.getvalue()

    def _insert_excel_ole(self, cell, evidence_dir: str):
        logger = get_logger()
        xlsx_path = os.path.join(evidence_dir, "package_dependency_list.xlsx")
        if not os.path.isfile(xlsx_path):
            return

        from lxml import etree
        import io
        from docx.opc.part import Part
        from docx.opc.packuri import PackURI

        icon_path = os.path.join(os.path.dirname(__file__), "..", "assets", "excel_icon.png")
        if not os.path.isfile(icon_path):
            logger.warning(f"Excel icon not found: {icon_path}")
            return

        with open(xlsx_path, "rb") as f:
            xlsx_bytes = f.read()
        with open(icon_path, "rb") as f:
            icon_bytes = f.read()

        ole_data = self._create_ole_raw_bin(xlsx_bytes)

        ole_para = cell.add_paragraph()
        ole_para.paragraph_format.space_before = Pt(4)
        ole_para.paragraph_format.space_after = Pt(4)

        ns_vml = "urn:schemas-microsoft-com:vml"
        ns_office = "urn:schemas-microsoft-com:office:office"
        ns_rel = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
        ns_word = "urn:schemas-microsoft-com:office:word"

        doc_part = cell.part
        pkg = doc_part.package

        ole_partname = PackURI("/word/embeddings/oleObject1.bin")
        ole_part = Part(ole_partname, "application/vnd.openxmlformats-officedocument.oleObject", ole_data, pkg)
        ole_rel = doc_part.relate_to(
            ole_part,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/oleObject",
        )

        icon_partname = PackURI("/word/media/excel_icon.png")
        icon_part = Part(icon_partname, "image/png", icon_bytes, pkg)
        icon_rel = doc_part.relate_to(
            icon_part,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/image",
        )

        ole_run = ole_para.add_run()

        shape_id = "_x0000_i1026"
        object_id = "_1468075725"

        obj_elem = etree.SubElement(_docx_element(ole_run), qn("w:object"))

        shape_elem = etree.SubElement(obj_elem, f"{{{ns_vml}}}shape")
        shape_elem.set("id", shape_id)
        shape_elem.set(f"{{{ns_office}}}spt", "75")
        shape_elem.set("type", "#_x0000_t75")
        shape_elem.set("style", "height:90pt;width:90pt;")
        shape_elem.set(f"{{{ns_office}}}ole", "t")
        shape_elem.set("filled", "f")
        shape_elem.set(f"{{{ns_office}}}preferrelative", "t")
        shape_elem.set("stroked", "f")
        shape_elem.set("coordsize", "21600,21600")

        path_elem = etree.SubElement(shape_elem, f"{{{ns_vml}}}path")
        fill_elem = etree.SubElement(shape_elem, f"{{{ns_vml}}}fill")
        fill_elem.set("on", "f")
        fill_elem.set("focussize", "0,0")

        stroke_elem = etree.SubElement(shape_elem, f"{{{ns_vml}}}stroke")
        stroke_elem.set("on", "f")

        imgdata_elem = etree.SubElement(shape_elem, f"{{{ns_vml}}}imagedata")
        imgdata_elem.set(f"{{{ns_rel}}}id", icon_rel)
        imgdata_elem.set(f"{{{ns_office}}}title", "oleimage")

        lock_elem = etree.SubElement(shape_elem, f"{{{ns_office}}}lock")
        lock_elem.set(f"{{{ns_vml}}}ext", "edit")
        lock_elem.set("aspectratio", "t")

        wrap_elem = etree.SubElement(shape_elem, f"{{{ns_word}}}wrap")
        wrap_elem.set("type", "none")
        anchor_elem = etree.SubElement(shape_elem, f"{{{ns_word}}}anchorlock")

        ole_obj_elem = etree.SubElement(obj_elem, f"{{{ns_office}}}OLEObject")
        ole_obj_elem.set("Type", "Embed")
        ole_obj_elem.set("ProgID", "Excel.Sheet.12")
        ole_obj_elem.set("ShapeID", shape_id)
        ole_obj_elem.set("DrawAspect", "Icon")
        ole_obj_elem.set("ObjectID", object_id)
        ole_obj_elem.set(f"{{{ns_rel}}}id", ole_rel)

        locked_elem = etree.SubElement(ole_obj_elem, f"{{{ns_office}}}LockedField")
        locked_elem.text = "false"

        logger.info(f"Excel OLE embedded: {os.path.basename(xlsx_path)}")

    @staticmethod
    def _embed_excel_attachments(docx_path: str, evidence_dir: str):
        logger = get_logger()
        if not os.path.isdir(evidence_dir):
            return
        xlsx_files = [f for f in os.listdir(evidence_dir) if f.endswith(".xlsx") and not f.startswith(("~$", ".~"))]
        log_files = [f for f in os.listdir(evidence_dir) if f.endswith("_log.txt")]
        attach_files = xlsx_files + log_files
        if not attach_files:
            return

        import zipfile
        import re as re_mod

        tmp_path = docx_path + ".tmp"

        with zipfile.ZipFile(docx_path, "r") as zin:
            with zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as zout:
                for item in zin.infolist():
                    data = zin.read(item.filename)
                    if item.filename == "[Content_Types].xml":
                        content = data.decode("utf-8")
                        if 'Extension="bin"' not in content:
                            bin_default = (
                                '<Default Extension="bin" '
                                'ContentType="application/vnd.openxmlformats'
                                '-officedocument.oleObject"/></Types>'
                            )
                            content = re_mod.sub(
                                '</Types>',
                                bin_default,
                                content,
                            )
                        if 'Extension="xlsx"' not in content:
                            xlsx_default = (
                                '<Default Extension="xlsx" '
                                'ContentType="application/vnd.openxmlformats'
                                '-officedocument.spreadsheetml.sheet"/></Types>'
                            )
                            content = re_mod.sub(
                                '</Types>',
                                xlsx_default,
                                content,
                            )
                        override_pattern = (
                            r'<Override\s+PartName="/word/embeddings/[^"]*"'
                            r'\s+ContentType="application/vnd\.openxmlformats'
                            r'-officedocument\.oleObject"\s*/>'
                        )
                        content = re_mod.sub(
                            override_pattern,
                            '',
                            content,
                        )
                        data = content.encode("utf-8")
                    zout.writestr(item, data)

                for xlsx_name in xlsx_files:
                    xlsx_path = os.path.join(evidence_dir, xlsx_name)
                    zout.write(xlsx_path, f"embeddings/{xlsx_name}")

                for log_name in log_files:
                    log_path = os.path.join(evidence_dir, log_name)
                    zout.write(log_path, f"embeddings/{log_name}")

        os.replace(tmp_path, docx_path)
        for name in attach_files:
            logger.info(f"File embedded in docx: {name}")

    def _insert_log_attachment(self, cell, log_path: str, test_label: str):
        logger = get_logger()
        log_name = os.path.basename(log_path)

        icon_path = os.path.join(os.path.dirname(__file__), "..", "assets", "log_icon.png")

        with open(log_path, "rb") as f:
            log_bytes = f.read()

        ole_data = self._create_ole_package_bin(log_bytes, filename=log_name)

        from lxml import etree
        from docx.opc.part import Part
        from docx.opc.packuri import PackURI

        ns_vml = "urn:schemas-microsoft-com:vml"
        ns_office = "urn:schemas-microsoft-com:office:office"
        ns_rel = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
        ns_word = "urn:schemas-microsoft-com:office:word"

        doc_part = cell.part
        pkg = doc_part.package

        ole_idx = getattr(self, '_ole_counter', 1) + 1
        self._ole_counter = ole_idx
        ole_partname = PackURI(f"/word/embeddings/oleObject{ole_idx}.bin")
        ole_part = Part(ole_partname, "application/vnd.openxmlformats-officedocument.oleObject", ole_data, pkg)
        ole_rel = doc_part.relate_to(
            ole_part,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/oleObject",
        )

        icon_rel = None
        if os.path.isfile(icon_path):
            with open(icon_path, "rb") as f:
                icon_bytes = f.read()
            icon_partname = PackURI(f"/word/media/log_icon_{ole_idx}.png")
            icon_part = Part(icon_partname, "image/png", icon_bytes, pkg)
            icon_rel = doc_part.relate_to(
                icon_part,
                "http://schemas.openxmlformats.org/officeDocument/2006/relationships/image",
            )

        ole_para = cell.add_paragraph()
        ole_para.paragraph_format.space_before = Pt(4)
        ole_para.paragraph_format.space_after = Pt(4)

        ole_run = ole_para.add_run()

        shape_id = f"_x0000_i102{ole_idx + 5}"
        object_id = f"_14680757{ole_idx + 25}"

        obj_elem = etree.SubElement(_docx_element(ole_run), qn("w:object"))

        shape_elem = etree.SubElement(obj_elem, f"{{{ns_vml}}}shape")
        shape_elem.set("id", shape_id)
        shape_elem.set(f"{{{ns_office}}}spt", "75")
        shape_elem.set("type", "#_x0000_t75")
        shape_elem.set("style", "height:90pt;width:90pt;")
        shape_elem.set(f"{{{ns_office}}}ole", "t")
        shape_elem.set("filled", "f")
        shape_elem.set(f"{{{ns_office}}}preferrelative", "t")
        shape_elem.set("stroked", "f")
        shape_elem.set("coordsize", "21600,21600")

        fill_elem = etree.SubElement(shape_elem, f"{{{ns_vml}}}fill")
        fill_elem.set("on", "f")
        fill_elem.set("focussize", "0,0")

        stroke_elem = etree.SubElement(shape_elem, f"{{{ns_vml}}}stroke")
        stroke_elem.set("on", "f")

        imgdata_elem = etree.SubElement(shape_elem, f"{{{ns_vml}}}imagedata")
        imgdata_elem.set(f"{{{ns_rel}}}id", icon_rel)
        imgdata_elem.set(f"{{{ns_office}}}title", "oleimage")

        lock_elem = etree.SubElement(shape_elem, f"{{{ns_office}}}lock")
        lock_elem.set(f"{{{ns_vml}}}ext", "edit")
        lock_elem.set("aspectratio", "t")

        wrap_elem = etree.SubElement(shape_elem, f"{{{ns_word}}}wrap")
        wrap_elem.set("type", "none")
        anchor_elem = etree.SubElement(shape_elem, f"{{{ns_word}}}anchorlock")

        ole_obj_elem = etree.SubElement(obj_elem, f"{{{ns_office}}}OLEObject")
        ole_obj_elem.set("Type", "Embed")
        ole_obj_elem.set("ProgID", "Package")
        ole_obj_elem.set("ShapeID", shape_id)
        ole_obj_elem.set("DrawAspect", "Icon")
        ole_obj_elem.set("ObjectID", object_id)
        ole_obj_elem.set(f"{{{ns_rel}}}id", ole_rel)

        locked_elem = etree.SubElement(ole_obj_elem, f"{{{ns_office}}}LockedField")
        locked_elem.text = "false"

        label_para = cell.add_paragraph()
        label_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        label_para.paragraph_format.space_before = Pt(0)
        label_para.paragraph_format.space_after = Pt(2)
        label_run = label_para.add_run(log_name)
        label_run.font.size = Pt(8)
        label_run.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)

        size_kb = os.path.getsize(log_path) / 1024
        line_count = 0
        try:
            with open(log_path, "r", encoding="utf-8", errors="replace") as f:
                line_count = sum(1 for _ in f)
        except Exception as e:
            logger.debug(f"Failed to count log lines: {e}")

        info_para = cell.add_paragraph()
        info_para.paragraph_format.space_before = Pt(0)
        info_para.paragraph_format.space_after = Pt(2)
        info_run = info_para.add_run(
            f"说明：{test_label} 测试完整执行日志（{line_count} 行，{size_kb:.1f} KB），"
            f"包含测试执行的命令输出、环境部署、用例运行等全部过程信息。双击图标可查看。"
        )
        info_run.font.size = Pt(8)
        info_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        logger.info(f"Log attachment noted in report: {log_name}")
